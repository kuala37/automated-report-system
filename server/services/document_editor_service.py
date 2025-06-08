import os
import json
from datetime import datetime
from docx import Document
from bs4 import BeautifulSoup
import mammoth
from models.models import Report, DocumentEdit
from sqlalchemy.ext.asyncio import AsyncSession
from docx.shared import Pt
from docx.enum.text import WD_UNDERLINE
from .docx_html_converter import WordToHtmlConverter

class DocumentEditorService:
    """Сервис для работы с документами в интерактивном режиме"""
    
    async def docx_to_html(self, docx_path):
        """Конвертация DOCX в HTML с использованием WordToHtmlConverter"""
        
        try:
            # Используем только новый конвертер
            converter = WordToHtmlConverter()
            enhanced_html = converter.convert_with_precise_formatting(docx_path)
            
            return enhanced_html
            
        except Exception as e:
            print(f"Ошибка при конвертации DOCX в HTML: {str(e)}")
            # Простой fallback без mammoth
            return await self._simple_fallback_conversion(docx_path)
    
    async def _simple_fallback_conversion(self, docx_path):
        """Простой fallback метод конвертации"""
        try:
            doc = Document(docx_path)
            html_parts = []
            
            html_parts.append('<!DOCTYPE html>')
            html_parts.append('<html><head><meta charset="utf-8">')
            html_parts.append('<style>')
            html_parts.append("""
                body { font-family: 'Times New Roman', serif; font-size: 12pt; padding: 20px; }
                .word-document-page { background: white; padding: 60px; margin: 0 auto; max-width: 21cm; }
                p { text-indent: 1.25cm; margin: 0; }
                h1, h2, h3, h4, h5, h6 { text-indent: 0; font-weight: bold; margin: 0; }
            """)
            html_parts.append('</style></head><body>')
            html_parts.append('<div class="word-document-page">')
            
            for i, paragraph in enumerate(doc.paragraphs):
                if paragraph.text.strip():
                    style_name = paragraph.style.name
                    if 'Heading' in style_name:
                        level = style_name.replace('Heading ', '')
                        html_parts.append(f'<h{level} data-paragraph-id="{i}">{paragraph.text}</h{level}>')
                    else:
                        html_parts.append(f'<p data-paragraph-id="{i}">{paragraph.text}</p>')
                else:
                    html_parts.append(f'<p data-paragraph-id="{i}" data-is-empty="true">&nbsp;</p>')
            
            html_parts.append('</div></body></html>')
            return '\n'.join(html_parts)
            
        except Exception as e:
            print(f"Fallback conversion error: {e}")
            return "<html><body><div class='word-document-page'><p>Ошибка конвертации документа</p></div></body></html>"
    
    async def update_document_with_edit(self, db: AsyncSession, report_id: int, edit_command: dict):
        """Обновляет документ на основе команды редактирования и создает новую версию"""
        # Получаем отчет
        report = await db.get(Report, report_id)
        if not report:
            raise ValueError(f"Отчет с ID {report_id} не найден")
        
        # Инициализируем version_history, если его нет
        if report.version_history is None:
            initial_version = {
                "version": 1,
                "timestamp": datetime.utcnow().isoformat(),
                "description": "Начальная версия документа",
                "file_path": report.file_path,
                "html_content": report.html_content,
                "edit_description": "Создание документа"
            }
            report.version_history = [initial_version]
            await db.commit()
        
        # Загружаем текущую версию документа
        doc = Document(report.file_path)
        
        # Выполняем команду
        command_type = edit_command.get("command")
        result = {"success": False, "message": "Неизвестная команда"}
        print(f"🔧 Выполняем команду: {edit_command}")
        
        # Устанавливаем переменные по умолчанию для избежания ошибок
        old_text = None
        new_text = None
        paragraph_id = None
        edit_description = f"Команда: {command_type}"
        
        if command_type == "replace_text":
            old_text = edit_command.get("oldText")
            new_text = edit_command.get("newText")
            paragraph_id = edit_command.get("paragraphId")
            result = await self._replace_text(doc, old_text, new_text, paragraph_id)
            edit_description = f"Замена текста: '{old_text[:30]}...' → '{new_text[:30]}...'"
            
        elif command_type == "format_text":
            text = edit_command.get("text")
            style = edit_command.get("style")
            paragraph_id = edit_command.get("paragraphId")
            result = await self._format_text(doc, text, style, paragraph_id)
            old_text = text
            new_text = f"{style}({text})"
            edit_description = f"Форматирование текста: {style}"
            
        elif command_type == "add_paragraph":
            text = edit_command.get("text")
            after_paragraph_id = edit_command.get("afterParagraphId")
            result = await self._add_paragraph(doc, text, after_paragraph_id)
            new_text = text
            edit_description = f"Добавление параграфа: '{text[:30]}...'"
            
        elif command_type == "delete_paragraph":
            paragraph_id = edit_command.get("paragraphId")
            result = await self._delete_paragraph(doc, paragraph_id)
            edit_description = f"Удаление параграфа {paragraph_id}"
            
        elif command_type == "rewrite_text":
            old_text = edit_command.get("oldText")
            new_text = edit_command.get("newText")
            paragraph_id = edit_command.get("paragraphId")
            result = await self._replace_text(doc, old_text, new_text, paragraph_id)
            edit_description = f"Перефразирование текста"
            
        elif command_type == "format_all_text":
            style = edit_command.get("style")
            result = await self._format_all_text(doc, style)
            old_text = "Весь текст документа"
            new_text = f"{style}(Весь текст документа)"
            edit_description = f"Форматирование всего текста: {style}"
            
        elif command_type == "format_all_headings":
            style = edit_command.get("style")
            result = await self._format_all_headings(doc, style)
            old_text = "Все заголовки документа"
            new_text = f"{style}(Все заголовки документа)"
            edit_description = f"Форматирование всех заголовков: {style}"
            
        elif command_type == "replace_all_occurrences":
            old_text = edit_command.get("oldText")
            new_text = edit_command.get("newText")
            result = await self._replace_all_text(doc, old_text, new_text)
            edit_description = f"Замена всех вхождений: '{old_text}' → '{new_text}'"
        
        elif command_type == "remove_formatting":
            text = edit_command.get("text")
            style = edit_command.get("style")
            paragraph_id = edit_command.get("paragraphId")
            print(f"Removing formatting: {text}, style: {style}, paragraph_id: {paragraph_id}")
            result = await self._remove_formatting(doc, text, style, paragraph_id)
            old_text = f"{style}({text})"
            new_text = text
            edit_description = f"Снятие форматирования {style}"
        
        elif command_type == "remove_all_formatting":
            result = await self._remove_all_formatting(doc)
            old_text = "Все форматирование документа"
            new_text = "Обычный текст"
            edit_description = "Снятие всего форматирования"

        # Если команда выполнена успешно, создаем новую версию
        if result["success"]:
            # Сохраняем текущую версию в истории (если еще не сохранена)
            version_history = list(report.version_history) if report.version_history else []
            current_version_exists = any(v.get("version") == report.document_version for v in version_history)
            
            if not current_version_exists:
                current_version_data = {
                    "version": report.document_version,
                    "timestamp": datetime.utcnow().isoformat(),
                    "description": f"Версия {report.document_version}",
                    "file_path": report.file_path,
                    "html_content": report.html_content,
                    "edit_description": "Предыдущее состояние"
                }
                version_history.append(current_version_data)
            
            # Создаем новую версию документа
            new_version = report.document_version + 1
            
            # Создаем путь для новой версии
            base_path, ext = os.path.splitext(report.file_path)
            # Убираем старый номер версии, если есть
            if base_path.endswith(f"_v{report.document_version}"):
                base_path = base_path[:-len(f"_v{report.document_version}")]
            new_file_path = f"{base_path}_v{new_version}{ext}"
            
            # Сохраняем обновленный документ
            doc.save(new_file_path)
            print(f"💾 Сохранена новая версия: {new_file_path}")
            
            # Конвертируем в HTML и сохраняем
            html_content = await self.docx_to_html(new_file_path)
            
            # Создаем запись для новой версии
            new_version_data = {
                "version": new_version,
                "timestamp": datetime.utcnow().isoformat(),
                "description": f"Версия {new_version}",
                "file_path": new_file_path,
                "html_content": html_content,
                "edit_description": edit_description
            }
            version_history.append(new_version_data)
            
            # Обновляем запись в базе данных
            report.document_version = new_version
            report.file_path = new_file_path
            report.html_content = html_content
            report.version_history = version_history
            
            # Сохраняем запись об изменении
            edit = DocumentEdit(
                report_id=report_id,
                user_id=edit_command.get("user_id"),
                chat_message_id=edit_command.get("message_id"),
                edit_type=command_type,
                content_before=json.dumps({"text": old_text}) if old_text else None,
                content_after=json.dumps({"text": new_text}) if new_text else None,
                position=json.dumps({"paragraph_id": paragraph_id}) if paragraph_id is not None else None
            )
            db.add(edit)
            await db.commit()
            
            print(f"✅ Создана версия {new_version}: {edit_description}")
        
        return result
    
    async def _replace_text(self, doc, old_text, new_text, paragraph_id=None):
        """Заменяет текст в документе"""
        if not old_text or not new_text:
            return {"success": False, "message": "Требуется указать старый и новый текст"}

        async def replace_with_boundaries_callback(paragraph, start_pos, end_pos, is_partial):
            """Заменяет точную часть параграфа"""
            
            full_text = paragraph.text
            
            # Очищаем параграф
            for i in range(len(paragraph.runs)):
                p = paragraph._p
                p.remove(paragraph.runs[0]._r)
            
            # Создаем новый текст с заменой
            new_full_text = full_text[:start_pos] + new_text + full_text[end_pos:]
            paragraph.add_run(new_full_text)
            
            return True

        # Пытаемся обработать как выделенный многопараграфный текст
        multi_result = await self._apply_operation_to_selected_text(
            doc, old_text, replace_with_boundaries_callback, "замены текста"
        )
        
        if multi_result is not None:
            return multi_result
        
        found = False
        
        if paragraph_id is not None:
            # Заменяем в конкретном параграфе
            try:
                paragraph = doc.paragraphs[int(paragraph_id)]
                if old_text in paragraph.text:
                    # Удаляем все runs и создаем новый с измененным текстом
                    text = paragraph.text.replace(old_text, new_text)
                    for i in range(len(paragraph.runs)):
                        p = paragraph._p
                        p.remove(paragraph.runs[0]._r)
                    paragraph.add_run(text)
                    found = True
            except IndexError:
                return {"success": False, "message": f"Параграф с ID {paragraph_id} не найден"}
        else:
            # Заменяем во всем документе
            for paragraph in doc.paragraphs:
                if old_text in paragraph.text:
                    # Удаляем все runs и создаем новый с измененным текстом
                    text = paragraph.text.replace(old_text, new_text)
                    for i in range(len(paragraph.runs)):
                        p = paragraph._p
                        p.remove(paragraph.runs[0]._r)
                    paragraph.add_run(text)
                    found = True
        
        if found:
            return {"success": True, "message": f"Текст '{old_text}' заменен на '{new_text}'"}
        else:
            return {"success": False, "message": f"Текст '{old_text}' не найден"}
    
    async def _format_text(self, doc, text, style, paragraph_id=None, start_pos=0, end_pos=None):
        """Форматирует текст в документе (поддержка многопараграфного текста)"""
        if not text or not style:
            return {"success": False, "message": "Требуется указать текст и стиль"}
        
        # Попробуем обработать как многопараграфный текст
        # Callback для форматирования с точными границами
        async def format_with_boundaries_callback(paragraph, start_pos, end_pos, is_partial):
            """Форматирует точную часть параграфа"""
            
            full_text = paragraph.text
            
            # Очищаем параграф
            for i in range(len(paragraph.runs)):
                p = paragraph._p
                p.remove(paragraph.runs[0]._r)
            
            # Создаем три части: до форматирования, форматированная часть, после форматирования
            if start_pos > 0:
                paragraph.add_run(full_text[:start_pos])
            
            # Форматированная часть
            if end_pos > start_pos:
                formatted_run = paragraph.add_run(full_text[start_pos:end_pos])
                if style == 'bold':
                    formatted_run.bold = True
                elif style == 'italic':
                    formatted_run.italic = True
                elif style == 'underline':
                    formatted_run.underline = WD_UNDERLINE.SINGLE
            
            # Часть после форматирования
            if end_pos < len(full_text):
                paragraph.add_run(full_text[end_pos:])
            
            return True
        
        # Пытаемся обработать как выделенный многопараграфный текст
        multi_result = await self._apply_operation_to_selected_text(
            doc, text, format_with_boundaries_callback, "форматирования"
        )
        
        if multi_result is not None:
            return multi_result
            
        # Если не многопараграфный текст, используем оригинальную логику для одного параграфа
        found = False
        
        if paragraph_id is not None:
            try:
                paragraph = doc.paragraphs[int(paragraph_id)]
                if text in paragraph.text:
                    # Находим позицию текста
                    if start_pos == 0 and end_pos is None:
                        start_pos = paragraph.text.find(text)
                        end_pos = start_pos + len(text)
                    elif end_pos is None:
                        end_pos = start_pos + len(text)
                    
                    # Очищаем параграф и создаем новые runs с форматированием
                    full_text = paragraph.text
                    for i in range(len(paragraph.runs)):
                        p = paragraph._p
                        p.remove(paragraph.runs[0]._r)
                    
                    # Создаем три части: до выделения, выделение, после выделения
                    if start_pos > 0:
                        paragraph.add_run(full_text[:start_pos])
                    
                    # Форматированный текст
                    formatted_run = paragraph.add_run(full_text[start_pos:end_pos])
                    if style == 'bold':
                        formatted_run.bold = True
                    elif style == 'italic':
                        formatted_run.italic = True
                    elif style == 'underline':
                        formatted_run.underline = WD_UNDERLINE.SINGLE
                    
                    # Текст после выделения
                    if end_pos < len(full_text):
                        paragraph.add_run(full_text[end_pos:])
                    
                    found = True
            except IndexError:
                return {"success": False, "message": f"Параграф с ID {paragraph_id} не найден"}
        else:
            # Поиск по всему документу
            for paragraph in doc.paragraphs:
                if text in paragraph.text:
                    start_pos = paragraph.text.find(text)
                    end_pos = start_pos + len(text)
                    
                    # Очищаем параграф и создаем новые runs с форматированием
                    full_text = paragraph.text
                    for i in range(len(paragraph.runs)):
                        p = paragraph._p
                        p.remove(paragraph.runs[0]._r)
                    
                    # Создаем три части
                    if start_pos > 0:
                        paragraph.add_run(full_text[:start_pos])
                    
                    formatted_run = paragraph.add_run(full_text[start_pos:end_pos])
                    if style == 'bold':
                        formatted_run.bold = True
                    elif style == 'italic':
                        formatted_run.italic = True
                    elif style == 'underline':
                        formatted_run.underline = WD_UNDERLINE.SINGLE
                    
                    if end_pos < len(full_text):
                        paragraph.add_run(full_text[end_pos:])
                    
                    found = True
        
        if found:
            return {"success": True, "message": f"Текст '{text[:50]}...' отформатирован как {style}"}
        else:
            return {"success": False, "message": f"Текст '{text[:50]}...' не найден"}

        
    async def _add_paragraph(self, doc, text, after_paragraph_id=None):
        """Добавляет новый параграф в документ"""
        if not text:
            return {"success": False, "message": "Требуется указать текст для нового параграфа"}
        
        try:
            if after_paragraph_id is not None:
                # Добавляем после указанного параграфа
                after_paragraph_id = int(after_paragraph_id)
                if after_paragraph_id < 0 or after_paragraph_id >= len(doc.paragraphs):
                    return {"success": False, "message": f"Параграф с ID {after_paragraph_id} не найден"}
                
                # Получаем элемент XML для вставки после указанного параграфа
                p = doc.paragraphs[after_paragraph_id]._p
                new_p = doc.add_paragraph(text)._p
                p.addnext(new_p)
            else:
                # Добавляем в конец документа
                doc.add_paragraph(text)
            
            return {"success": True, "message": f"Добавлен новый параграф: '{text[:30]}...'"}
        except Exception as e:
            return {"success": False, "message": f"Ошибка при добавлении параграфа: {str(e)}"}

    async def _format_all_text(self, doc, style):
        """Применяет форматирование ко всему тексту в документе"""
        if style not in ['bold', 'italic', 'underline']:
            return {"success": False, "message": f"Неподдерживаемый стиль форматирования: {style}"}
        
        count = 0
        # Обрабатываем все параграфы
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():  # Проверяем, что параграф не пустой
                # Применяем форматирование ко всему параграфу
                for run in paragraph.runs:
                    if style == 'bold':
                        run.bold = True
                    elif style == 'italic':
                        run.italic = True
                    elif style == 'underline':
                        run.underline = WD_UNDERLINE.SINGLE
                
                if not paragraph.runs:
                    run = paragraph.add_run(paragraph.text)
                    if style == 'bold':
                        run.bold = True
                    elif style == 'italic':
                        run.italic = True
                    elif style == 'underline':
                        run.underline = WD_UNDERLINE.SINGLE
                
                count += 1
        
        style_name_map = {
            'bold': 'жирным',
            'italic': 'курсивом',
            'underline': 'подчеркнутым'
        }
        return {"success": True, "message": f"Весь текст в документе сделан {style_name_map.get(style, style)} ({count} параграфов)"}

    async def _format_all_headings(self, doc, style):
        """Форматирует все заголовки в документе"""
        if style not in ['bold', 'italic', 'underline']:
            return {"success": False, "message": f"Неподдерживаемый стиль форматирования: {style}"}
        
        count = 0
        # Находим все параграфы, которые являются заголовками
        for paragraph in doc.paragraphs:
            if paragraph.style.name.startswith('Heading'):
                # Применяем форматирование ко всему заголовку
                for run in paragraph.runs:
                    if style == 'bold':
                        run.bold = True
                    elif style == 'italic':
                        run.italic = True
                    elif style == 'underline':
                        run.underline = WD_UNDERLINE.SINGLE
                
                # Если в заголовке нет runs, создаем новый
                if not paragraph.runs:
                    run = paragraph.add_run(paragraph.text)
                    if style == 'bold':
                        run.bold = True
                    elif style == 'italic':
                        run.italic = True
                    elif style == 'underline':
                        run.underline = WD_UNDERLINE.SINGLE
                
                count += 1
        
        style_name_map = {
            'bold': 'жирным',
            'italic': 'курсивом',
            'underline': 'подчеркнутым'
        }
        return {"success": True, "message": f"Все заголовки сделаны {style_name_map.get(style, style)} ({count} заголовков)"}

    async def _replace_all_text(self, doc, old_text, new_text):
        """Заменяет все вхождения текста во всем документе"""
        if not old_text or not new_text:
            return {"success": False, "message": "Требуется указать старый и новый текст"}
        
        count = 0
        # Обрабатываем все параграфы
        for paragraph in doc.paragraphs:
            if old_text in paragraph.text:
                # Заменяем текст в параграфе
                new_paragraph_text = paragraph.text.replace(old_text, new_text)
                
                # Очищаем все runs
                for i in range(len(paragraph.runs)):
                    p = paragraph._p
                    p.remove(paragraph.runs[0]._r)
                
                # Добавляем новый run с обновленным текстом
                paragraph.add_run(new_paragraph_text)
                count += 1
        
        if count > 0:
            return {"success": True, "message": f"Текст '{old_text}' заменен на '{new_text}' {count} раз по всему документу"}
        else:
            return {"success": False, "message": f"Текст '{old_text}' не найден в документе"}    
    
    async def _delete_paragraph(self, doc, paragraph_id):
        """Удаляет параграф из документа"""
        try:
            paragraph_id = int(paragraph_id)
            
            # Создаем маппинг видимых параграфов
            visible_to_docx, docx_to_visible = await self._get_visible_paragraph_mapping(doc)
            
            print(f"Маппинг видимых параграфов: {visible_to_docx}")
            print(f"Пользователь просит удалить параграф ID: {paragraph_id}")
            
            # Определяем реальный номер параграфа в DOCX
            actual_docx_id = None
            
            # Если paragraph_id меньше количества видимых параграфов, 
            # значит пользователь имеет в виду видимый параграф
            if paragraph_id < len(visible_to_docx):
                actual_docx_id = visible_to_docx[paragraph_id]
                print(f"Преобразуем видимый параграф {paragraph_id} в DOCX параграф {actual_docx_id}")
            else:
                # Иначе считаем, что это уже DOCX ID
                actual_docx_id = paragraph_id
                print(f"Используем DOCX ID напрямую: {actual_docx_id}")
            
            if actual_docx_id < 0 or actual_docx_id >= len(doc.paragraphs):
                return {"success": False, "message": f"Параграф с ID {paragraph_id} не найден. Всего параграфов: {len(doc.paragraphs)}"}
            
            # Получаем текст для сообщения
            target_paragraph = doc.paragraphs[actual_docx_id]
            paragraph_text = target_paragraph.text[:30] if target_paragraph.text.strip() else "[ПУСТОЙ ПАРАГРАФ]"
            
            print(f"Удаляем параграф {actual_docx_id}: '{paragraph_text}'")
            
            # Удаляем параграф
            p = target_paragraph._p
            p.getparent().remove(p)
            
            return {"success": True, "message": f"Удален параграф: '{paragraph_text}...'"}
        except Exception as e:
            print(f"Ошибка при удалении параграфа: {e}")
            return {"success": False, "message": f"Ошибка при удалении параграфа: {str(e)}"}

    async def _remove_formatting(self, doc, text, style, paragraph_id=None):
        """Снимает конкретное форматирование с текста (поддержка многопараграфного)"""
        if not text:
            return {"success": False, "message": "Требуется указать текст"}
        
        # Попробуем обработать как многопараграфный текст
        async def remove_format_with_boundaries_callback(paragraph, start_pos, end_pos, is_partial):
            """Снимает форматирование с точной части параграфа"""
            
            # Сохраняем оригинальное форматирование всех runs
            original_runs_data = []
            for run in paragraph.runs:
                original_runs_data.append({
                    'text': run.text,
                    'bold': run.bold,
                    'italic': run.italic,
                    'underline': run.underline
                })
            
            # Очищаем параграф
            full_text = paragraph.text
            for i in range(len(paragraph.runs)):
                p = paragraph._p
                p.remove(paragraph.runs[0]._r)
            
            # Восстанавливаем текст с правильным форматированием
            current_pos = 0
            
            for run_data in original_runs_data:
                run_text = run_data['text']
                run_start = current_pos
                run_end = current_pos + len(run_text)
                
                # Определяем пересечение с выделенным текстом
                overlap_start = max(run_start, start_pos)
                overlap_end = min(run_end, end_pos)
                
                if overlap_start < overlap_end:
                    # Есть пересечение - разделяем run
                    
                    # Часть до пересечения
                    if overlap_start > run_start:
                        before_run = paragraph.add_run(run_text[:overlap_start - run_start])
                        before_run.bold = run_data['bold']
                        before_run.italic = run_data['italic']
                        before_run.underline = run_data['underline']
                    
                    # Пересекающаяся часть (снимаем форматирование)
                    overlap_text = run_text[overlap_start - run_start:overlap_end - run_start]
                    unformatted_run = paragraph.add_run(overlap_text)
                    if style != 'bold':
                        unformatted_run.bold = run_data['bold']
                    if style != 'italic':
                        unformatted_run.italic = run_data['italic']
                    if style != 'underline':
                        unformatted_run.underline = run_data['underline']
                    
                    # Часть после пересечения
                    if overlap_end < run_end:
                        after_run = paragraph.add_run(run_text[overlap_end - run_start:])
                        after_run.bold = run_data['bold']
                        after_run.italic = run_data['italic']
                        after_run.underline = run_data['underline']
                else:
                    # Нет пересечения - восстанавливаем как есть
                    restored_run = paragraph.add_run(run_text)
                    restored_run.bold = run_data['bold']
                    restored_run.italic = run_data['italic']
                    restored_run.underline = run_data['underline']
                
                current_pos = run_end
            
            return True

        # Пытаемся обработать как выделенный многопараграфный текст
        multi_result = await self._apply_operation_to_selected_text(
            doc, text, remove_format_with_boundaries_callback, "снятия форматирования"
        )
        
        if multi_result is not None:
            style_name_map = {
                'bold': 'жирности',
                'italic': 'курсива', 
                'underline': 'подчеркивания'
            }
            
            if style in style_name_map:
                return {"success": True, "message": f"Снято форматирование {style_name_map[style]} с выделенного текста"}
            else:
                return {"success": True, "message": f"Снято все форматирование с выделенного текста"}

        # Оригинальная логика для одного параграфа (оставляем существующую)
        found = False
        actual_paragraph_id = None
        
        # Сначала ищем текст по всему документу
        print(f"🔍 Ищем текст '{text}' в документе...")
        for i, paragraph in enumerate(doc.paragraphs):
            if text in paragraph.text:
                found = True
                actual_paragraph_id = i
                print(f"✅ Текст найден в параграфе {i}")
                break
        
        if not found:
            return {"success": False, "message": f"Текст '{text}' не найден в документе"}
        
        # Если указан paragraph_id, но текст найден в другом параграфе - предупреждаем
        if paragraph_id is not None and int(paragraph_id) != actual_paragraph_id:
            print(f"⚠️ ВНИМАНИЕ: Frontend передал paragraph_id={paragraph_id}, но текст найден в параграфе {actual_paragraph_id}")
        
        # Работаем с правильным параграфом
        try:
            paragraph = doc.paragraphs[actual_paragraph_id]
            print(f"Обрабатываем параграф {actual_paragraph_id}")
            
            # Находим позицию искомого текста
            start_pos = paragraph.text.find(text)
            end_pos = start_pos + len(text)
            
            print(f"Позиция текста: {start_pos}-{end_pos}")
            
            # Сохраняем форматирование каждого run
            original_runs_data = []
            for run in paragraph.runs:
                original_runs_data.append({
                    'text': run.text,
                    'bold': run.bold,
                    'italic': run.italic,
                    'underline': run.underline
                })
            
            # Очищаем параграф
            full_text = paragraph.text
            for i in range(len(paragraph.runs)):
                p = paragraph._p
                p.remove(paragraph.runs[0]._r)
            
            # Восстанавливаем текст с правильным форматированием
            current_pos = 0
            
            for run_data in original_runs_data:
                run_text = run_data['text']
                run_start = current_pos
                run_end = current_pos + len(run_text)
                
                # Определяем пересечение с выделенным текстом
                overlap_start = max(run_start, start_pos)
                overlap_end = min(run_end, end_pos)
                
                if overlap_start < overlap_end:
                    # Есть пересечение - нужно разделить run
                    
                    # Часть до пересечения (сохраняем оригинальное форматирование)
                    if overlap_start > run_start:
                        before_run = paragraph.add_run(run_text[:overlap_start - run_start])
                        before_run.bold = run_data['bold']
                        before_run.italic = run_data['italic']
                        before_run.underline = run_data['underline']
                    
                    # Пересекающаяся часть (снимаем форматирование)
                    overlap_text = run_text[overlap_start - run_start:overlap_end - run_start]
                    unformatted_run = paragraph.add_run(overlap_text)
                    unformatted_run.bold = False
                    unformatted_run.italic = False
                    unformatted_run.underline = False
                    
                    # Часть после пересечения (сохраняем оригинальное форматирование)
                    if overlap_end < run_end:
                        after_run = paragraph.add_run(run_text[overlap_end - run_start:])
                        after_run.bold = run_data['bold']
                        after_run.italic = run_data['italic']
                        after_run.underline = run_data['underline']
                else:
                    # Нет пересечения - просто восстанавливаем run как есть
                    restored_run = paragraph.add_run(run_text)
                    restored_run.bold = run_data['bold']
                    restored_run.italic = run_data['italic']
                    restored_run.underline = run_data['underline']
                
                current_pos = run_end
            
            style_name_map = {
                'bold': 'жирности',
                'italic': 'курсива', 
                'underline': 'подчеркивания'
            }
            
            if style in style_name_map:
                return {"success": True, "message": f"Снято форматирование {style_name_map[style]} с текста '{text}' в параграфе {actual_paragraph_id}"}
            else:
                return {"success": True, "message": f"Снято все форматирование с текста '{text}' в параграфе {actual_paragraph_id}"}
                    
        except Exception as e:
            print(f"Ошибка при обработке параграфа: {e}")
            return {"success": False, "message": f"Ошибка при обработке: {str(e)}"}

    async def _remove_all_formatting(self, doc):
        """Снимает все форматирование с конкретного текста"""
        try:
            count = 0
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():  
                    for run in paragraph.runs:
                        run.bold = False
                        run.italic = False
                        run.underline = False
                    count += 1
            return {"success": True, "message": f"Снято все форматирование с документа ({count} параграфов)"}
        except Exception as e:
            return {"success": False, "message": f"Ошибка удалении форматирования со всего текста: {str(e)}"}


    # async def _remove_all_formatting_document(self, doc):
    #     """Снимает все форматирование со всего документа"""
    #     count = 0
        
    #     for paragraph in doc.paragraphs:
    #         if paragraph.text.strip():  # Проверяем, что параграф не пустой
    #             # Снимаем форматирование со всех runs
    #             for run in paragraph.runs:
    #                 run.bold = False
    #                 run.italic = False
    #                 run.underline = False
    #             count += 1
        
    #     return {"success": True, "message": f"Снято все форматирование с документа ({count} параграфов)"}
    


    async def _get_visible_paragraph_mapping(self, doc):
        """Создает маппинг между видимыми параграфами и реальными номерами в DOCX"""
        visible_to_docx = {}
        docx_to_visible = {}
        visible_index = 0
        
        for docx_index, paragraph in enumerate(doc.paragraphs):
            if paragraph.text.strip():  # Только непустые параграфы видны в HTML
                visible_to_docx[visible_index] = docx_index
                docx_to_visible[docx_index] = visible_index
                visible_index += 1
        
        return visible_to_docx, docx_to_visible


    async def _convert_user_paragraph_number_to_docx(self, doc, user_paragraph_number):
        """Преобразует номер параграфа от пользователя (видимый) в реальный номер DOCX"""
        visible_to_docx, _ = await self._get_visible_paragraph_mapping(doc)
        
        # Пользователь считает с 1, а мы с 0
        visible_index = user_paragraph_number - 1
        
        if visible_index in visible_to_docx:
            return visible_to_docx[visible_index]
        else:
            return None
        

    def _is_multi_paragraph_text(self, text):
        """Проверяет, является ли текст многопараграфным"""
        return len(text) > 200 or '\n' in text or '"' in text or '*' in text
        #return '\n' in text or '"' in text or '*' in text

    def _extract_text_parts(self, text):
        """Извлекает части текста для поиска в документе"""
        normalized_text = ' '.join(text.split())
        text_parts = []
        
        # Разбиваем по специальным символам (списки, заголовки)
        if '*' in normalized_text:
            parts = [part.strip() for part in normalized_text.split('*') if part.strip() and len(part.strip()) > 5]
            text_parts.extend(parts)
        
        # Разбиваем по знакам препинания
        for delimiter in ['.', ':', '###', '1.', '2.', '3.', '4.', '5.']:
            if delimiter in normalized_text:
                parts = [part.strip() for part in normalized_text.split(delimiter) if part.strip() and len(part.strip()) > 10]
                text_parts.extend(parts)
        
        # Если ничего не нашли, разбиваем по длине
        if not text_parts:
            words = normalized_text.split()
            current_part = ""
            for word in words:
                if len(current_part + " " + word) < 80:
                    current_part += " " + word if current_part else word
                else:
                    if current_part and len(current_part) > 10:
                        text_parts.append(current_part.strip())
                    current_part = word
            if current_part and len(current_part) > 10:
                text_parts.append(current_part.strip())
        
        # Удаляем дубликаты и сортируем по длине (длинные части сначала)
        text_parts = list(set(text_parts))
        text_parts.sort(key=len, reverse=True)
        
        return text_parts, normalized_text

    def _find_matching_paragraphs(self, doc, text_parts, operation_name="обработки"):
        """Находит параграфы, соответствующие частям текста"""
        matched_paragraphs = []
        processed_paragraphs = set()
        
        print(f"📝 Разбили текст на {len(text_parts)} частей для {operation_name}")
        for i, part in enumerate(text_parts[:5]):  # Показываем первые 5 частей
            print(f"  {i+1}. '{part[:50]}...' (длина: {len(part)})")
        
        # Первый проход: вычисляем "вес" каждого параграфа (сколько частей текста содержит)
        paragraph_scores = {}
        paragraph_matched_parts = {}
        
        for paragraph_idx, paragraph in enumerate(doc.paragraphs):
            paragraph_text = paragraph.text.strip()
            if not paragraph_text:
                continue
            
            normalized_paragraph = ' '.join(paragraph_text.split())
            matched_parts_for_paragraph = []
            score = 0
            
            # Проверяем ВСЕ части текста для каждого параграфа
            for part in text_parts:
                if len(part) < 10:  # Слишком короткие части игнорируем
                    continue
                
                part_matched = False
                
                # Точное совпадение
                if part in normalized_paragraph:
                    matched_parts_for_paragraph.append({
                        'part': part,
                        'match_type': 'exact',
                        'score': len(part) / len(normalized_paragraph)
                    })
                    score += 2 * len(part)  # Больший вес для точных совпадений
                    part_matched = True
                
                # Если нет точного совпадения, проверяем частичное и фразовое
                if not part_matched:
                    # Частичное совпадение (общие слова)
                    part_words = set(part.lower().split())
                    paragraph_words = set(normalized_paragraph.lower().split())
                    common_words = part_words.intersection(paragraph_words)
                    
                    word_ratio = len(common_words) / len(part_words) if part_words else 0
                    
                    if word_ratio >= 0.7 and len(part_words) >= 3:
                        matched_parts_for_paragraph.append({
                            'part': part,
                            'match_type': 'partial',
                            'score': word_ratio,
                            'common_words': len(common_words),
                            'total_words': len(part_words)
                        })
                        score += word_ratio * len(part)
                        part_matched = True
                    
                    # Ключевые фразы (3+ слова подряд)
                    if not part_matched and len(part.split()) >= 3:
                        part_words_list = part.split()
                        for i in range(len(part_words_list) - 2):
                            phrase = ' '.join(part_words_list[i:i+3])
                            if phrase.lower() in normalized_paragraph.lower():
                                matched_parts_for_paragraph.append({
                                    'part': part,
                                    'match_type': 'phrase',
                                    'score': 3 / len(part_words_list),
                                    'phrase': phrase
                                })
                                score += 0.5 * len(phrase)
                                part_matched = True
                                break
            
            # Сохраняем результаты только если есть совпадения
            if matched_parts_for_paragraph:
                paragraph_scores[paragraph_idx] = score
                paragraph_matched_parts[paragraph_idx] = matched_parts_for_paragraph
        
        # Сортируем параграфы по весу (от большего к меньшему)
        sorted_paragraphs = sorted(
            paragraph_scores.items(), 
            key=lambda x: x[1], 
            reverse=True
        )
        
        # Выводим информацию о лучших совпадениях
        print(f"🔍 Найдены следующие совпадающие параграфы (по весу):")
        for idx, (p_idx, score) in enumerate(sorted_paragraphs[:6]):
            paragraph = doc.paragraphs[p_idx]
            print(f"  {idx+1}. Параграф {p_idx} (вес: {score:.2f}): '{paragraph.text[:50]}...'")
            
            for match in paragraph_matched_parts[p_idx][:3]:  # Top 3 matches
                match_type = match['match_type']
                if match_type == 'exact':
                    print(f"     ✓ Точное совпадение: '{match['part'][:40]}...'")
                elif match_type == 'partial':
                    print(f"     ≈ Частичное совпадение: {match['common_words']}/{match['total_words']} слов")
                elif match_type == 'phrase':
                    print(f"     ~ Фразовое совпадение: '{match['phrase']}'")
        
        # Добавляем параграфы с высоким весом
        threshold = max(paragraph_scores.values()) * 0.3 if paragraph_scores else 0
        
        for p_idx, score in sorted_paragraphs:
            if p_idx in processed_paragraphs:
                continue
                
            # Принимаем параграфы с весом выше порога
            if score >= threshold:
                paragraph = doc.paragraphs[p_idx]
                matched_parts = [m['part'] for m in paragraph_matched_parts[p_idx]]
                
                print(f"✅ Выбран параграф {p_idx} (вес: {score:.2f}): '{paragraph.text[:50]}...'")
                
                matched_paragraphs.append({
                    'index': p_idx,
                    'paragraph': paragraph,
                    'matched_parts': matched_parts,
                    'score': score
                })
                processed_paragraphs.add(p_idx)
        
        # Если не нашли никаких параграфов, возвращаем пустой список
        if not matched_paragraphs:
            print("⚠️ Не найдены параграфы, соответствующие частям текста")
            
        # Сортируем выбранные параграфы по их индексу в документе
        matched_paragraphs.sort(key=lambda x: x['index'])
        
        return matched_paragraphs


    def _fallback_search_by_keywords(self, doc, normalized_text, operation_name="обработки"):
        """Альтернативный поиск по ключевым словам"""
        matched_paragraphs = []
        unique_words = [word for word in normalized_text.split() if len(word) > 4 and not word.isdigit()][:10]
        print(f"🔍 Альтернативный поиск для {operation_name} по ключевым словам: {unique_words[:5]}...")
        
        for paragraph_idx, paragraph in enumerate(doc.paragraphs):
            paragraph_text = paragraph.text.strip()
            if not paragraph_text:
                continue
            
            paragraph_words = paragraph_text.lower().split()
            matches = sum(1 for word in unique_words if word.lower() in paragraph_words)
            
            if matches >= 2:  # Если найдено 2+ ключевых слова
                print(f"✅ Альтернативное совпадение в параграфе {paragraph_idx}: {matches} ключевых слов")
                matched_paragraphs.append({
                    'index': paragraph_idx,
                    'paragraph': paragraph,
                    'matched_parts': [f"keywords: {matches} matches"]
                })
        
        return matched_paragraphs

    async def _process_multi_paragraph_text(self, doc, text, operation_callback, operation_name="обработки"):
        """Универсальный метод для обработки многопараграфного текста"""
        if not self._is_multi_paragraph_text(text):
            return None  # Не многопараграфный текст
        
        print(f"🔍 Обнаружен многопараграфный текст для {operation_name}, длина: {len(text)}")
        
        # Извлекаем части текста
        text_parts, normalized_text = self._extract_text_parts(text)
        
        # Ищем соответствующие параграфы
        matched_paragraphs = self._find_matching_paragraphs(doc, text_parts, operation_name)
        
        # Если ничего не найдено, пробуем альтернативный поиск
        if not matched_paragraphs:
            print("❌ Основной поиск не дал результатов")
            matched_paragraphs = self._fallback_search_by_keywords(doc, normalized_text, operation_name)
        
        if not matched_paragraphs:
            return {"success": False, "message": f"Многопараграфный текст не найден в документе"}
        
        # Обрабатываем найденные параграфы
        processed_count = 0
        for match in matched_paragraphs:
            paragraph_idx = match['index']
            paragraph = match['paragraph']
            matched_parts = match['matched_parts']
            
            print(f"✅ Обрабатываем параграф {paragraph_idx}: '{paragraph.text[:50]}...'")
            print(f"   Совпавшие части: {[str(p)[:30] + '...' if len(str(p)) > 30 else str(p) for p in matched_parts]}")
            
            # Вызываем callback для обработки конкретного параграфа
            if await operation_callback(paragraph):
                processed_count += 1
        
        return {
            "success": True, 
            "message": f"Обработано {processed_count} параграфов ({operation_name})",
            "processed_count": processed_count
        }
    

    def _find_text_boundaries_in_paragraphs(self, doc, selected_text, matched_paragraphs):
        """Находит точные границы выделенного текста в найденных параграфах"""
        text_boundaries = []
        remaining_text = selected_text.strip()
        
        print(f"🎯 Определяем границы для текста: '{selected_text[:100]}...'")
        
        for idx, match in enumerate(matched_paragraphs):
            paragraph = match['paragraph']
            paragraph_text = paragraph.text
            
            boundary_info = {
                'paragraph_index': match['index'],
                'paragraph': paragraph,
                'start_pos': 0,
                'end_pos': len(paragraph_text),
                'is_partial': False
            }
            
            # Первый параграф - ищем начало выделения
            if idx == 0 and len(matched_paragraphs) > 1:
                # Ищем где начинается выделенный текст в первом параграфе
                start_words = selected_text.split()[:5]  # Первые 5 слов
                start_phrase = ' '.join(start_words)
                
                start_pos = paragraph_text.find(start_phrase)
                if start_pos >= 0:
                    boundary_info['start_pos'] = start_pos
                    boundary_info['is_partial'] = True
                    print(f"📍 Первый параграф: начало с позиции {start_pos}")
                else:
                    # Поиск по словам
                    for word in start_words:
                        if word in paragraph_text:
                            start_pos = paragraph_text.find(word)
                            boundary_info['start_pos'] = start_pos
                            boundary_info['is_partial'] = True
                            print(f"📍 Первый параграф: начало найдено по слову '{word}' на позиции {start_pos}")
                            break
            
            # Последний параграф - ищем конец выделения
            elif idx == len(matched_paragraphs) - 1 and len(matched_paragraphs) > 1:
                # Ищем где заканчивается выделенный текст
                end_words = selected_text.split()[-5:]  # Последние 5 слов
                end_phrase = ' '.join(end_words)
                
                # Убираем знаки препинания для более точного поиска
                clean_end_phrase = end_phrase.rstrip('.,!?;: ')
                
                end_pos = paragraph_text.find(clean_end_phrase)
                if end_pos >= 0:
                    boundary_info['end_pos'] = end_pos + len(clean_end_phrase)
                    boundary_info['is_partial'] = True
                    print(f"📍 Последний параграф: конец на позиции {boundary_info['end_pos']}")
                else:
                    # Поиск по последним словам
                    for word in reversed(end_words):
                        word_pos = paragraph_text.find(word)
                        if word_pos >= 0:
                            boundary_info['end_pos'] = word_pos + len(word)
                            boundary_info['is_partial'] = True
                            print(f"📍 Последний параграф: конец найден по слову '{word}' на позиции {boundary_info['end_pos']}")
                            break
            
            # Единственный параграф - ищем точные границы выделения
            elif len(matched_paragraphs) == 1:
                # Находим точное вхождение текста
                clean_selected = ' '.join(selected_text.split())
                clean_paragraph = ' '.join(paragraph_text.split())
                
                start_pos = clean_paragraph.find(clean_selected)
                if start_pos >= 0:
                    # Переводим позицию обратно в оригинальный текст
                    original_start = self._find_original_position(paragraph_text, start_pos, clean_paragraph)
                    original_end = original_start + len(selected_text)
                    
                    boundary_info['start_pos'] = original_start
                    boundary_info['end_pos'] = min(original_end, len(paragraph_text))
                    boundary_info['is_partial'] = True
                    print(f"📍 Единственный параграф: позиции {original_start}-{boundary_info['end_pos']}")
            
            text_boundaries.append(boundary_info)
        
        return text_boundaries

    def _find_original_position(self, original_text, normalized_pos, normalized_text):
        """Находит позицию в оригинальном тексте по позиции в нормализованном"""
        if normalized_pos == 0:
            return 0
        
        original_words = original_text.split()
        normalized_words = normalized_text.split()
        
        # Подсчитываем количество символов до нужной позиции
        char_count = 0
        word_count = 0
        
        for char in normalized_text[:normalized_pos]:
            if char == ' ':
                word_count += 1
            char_count += 1
        
        # Находим соответствующую позицию в оригинальном тексте
        original_pos = 0
        current_word = 0
        
        for i, char in enumerate(original_text):
            if char == ' ' or char == '\t' or char == '\n':
                if current_word < word_count:
                    current_word += 1
                original_pos = i
            elif current_word >= word_count:
                return original_pos
            
            if current_word >= word_count:
                return i
        
        return min(original_pos, len(original_text))
    

    async def _apply_operation_to_selected_text(self, doc, selected_text, operation_callback, operation_name="обработки"):
        """Универсальный метод для применения операций к точно выделенному тексту"""
        
        # Проверяем, является ли текст многопараграфным
        if not self._is_multi_paragraph_text(selected_text):
            return None  # Обрабатывается стандартным способом
        
        print(f"🎯 Применяем {operation_name} к выделенному тексту длиной {len(selected_text)}")
        
        # Находим соответствующие параграфы
        text_parts, normalized_text = self._extract_text_parts(selected_text)
        matched_paragraphs = self._find_matching_paragraphs(doc, text_parts, operation_name)
        
        if not matched_paragraphs:
            matched_paragraphs = self._fallback_search_by_keywords(doc, normalized_text, operation_name)
        
        if not matched_paragraphs:
            return {"success": False, "message": f"Выделенный текст не найден в документе"}
        
        # Определяем точные границы выделения в каждом параграфе
        text_boundaries = self._find_text_boundaries_in_paragraphs(doc, selected_text, matched_paragraphs)
        
        # Применяем операцию к каждому параграфу с учетом границ
        processed_count = 0
        for boundary in text_boundaries:
            paragraph = boundary['paragraph']
            start_pos = boundary['start_pos']
            end_pos = boundary['end_pos']
            is_partial = boundary['is_partial']
            
            print(f"✅ Обрабатываем параграф {boundary['paragraph_index']}: позиции {start_pos}-{end_pos}")
            print(f"   Обрабатываемый текст: '{paragraph.text[start_pos:end_pos]}'")
            
            # Вызываем callback с информацией о границах
            if await operation_callback(paragraph, start_pos, end_pos, is_partial):
                processed_count += 1
        
        return {
            "success": True,
            "message": f"Выполнено {operation_name} для {processed_count} параграфов",
            "processed_count": processed_count
        }

