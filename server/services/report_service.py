import traceback
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
import os
from generation.generate_text_langchain import generate_text, generate_text_with_params

class GostStyle:
    """GOST document styling configuration"""
    def __init__(self, gost_type="8.5"):
        self.gost_type = gost_type
        # ГОСТ 7.32-2017 
        if gost_type == "7.32":
            self.margins = {
                "top": Cm(2), "bottom": Cm(2),
                "left": Cm(2.5), "right": Cm(1.5)
            }
            self.font_name = "Times New Roman"
            self.main_font_size = Pt(14)
            self.heading_font_size = Pt(16)
            self.line_spacing = 1.5
        # ГОСТ 8.5 
        else:
            self.margins = {
                "top": Cm(2), "bottom": Cm(2),
                "left": Cm(3), "right": Cm(1.5)
            }
            self.font_name = "Times New Roman"
            self.main_font_size = Pt(12)
            self.heading_font_size = Pt(14)
            self.line_spacing = 1.5

class WordDocumentGenerator:
    def __init__(self, gost_style=None):
        self.document = Document()
        self.style = gost_style if gost_style else GostStyle()
        self._apply_gost_formatting()

    def _apply_gost_formatting(self):
        """Apply GOST formatting to the document"""
        section = self.document.sections[0]
        # Set margins
        section.top_margin = self.style.margins["top"]
        section.bottom_margin = self.style.margins["bottom"]
        section.left_margin = self.style.margins["left"]
        section.right_margin = self.style.margins["right"]

        # Set default paragraph style
        style = self.document.styles['Normal']
        font = style.font
        font.name = self.style.font_name
        font.size = self.style.main_font_size
        paragraph_format = style.paragraph_format
        paragraph_format.line_spacing = self.style.line_spacing

    def add_title(self, title, size=None):
        """Add a title to the document with specified size"""
        heading = self.document.add_heading(title, 0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in heading.runs:
            run.font.name = self.style.font_name
            run.font.size = size if size else self.style.heading_font_size

    def add_section(self, title, content, heading_level=1):
        """Add a section with heading and content"""
        heading = self.document.add_heading(title, level=heading_level)
        for run in heading.runs:
            run.font.name = self.style.font_name
            run.font.size = self.style.heading_font_size

        paragraph = self.document.add_paragraph(content)
        for run in paragraph.runs:
            run.font.name = self.style.font_name
            run.font.size = self.style.main_font_size

    def add_generated_section(self, prompt, section_title, heading_level=1, **generation_params):
        """Add a section with AI-generated content"""
        self.document.add_heading(section_title, level=heading_level)
        
        generated_content = generate_text_with_params(prompt, **generation_params)
        
        self.document.add_paragraph(generated_content)
        
        return generated_content

    def add_paragraph_text(self, text):
        """Add a simple paragraph of text"""
        paragraph = self.document.add_paragraph(text)
        for run in paragraph.runs:
            run.font.name = self.style.font_name
            run.font.size = self.style.main_font_size

    def save(self, filename):
        """Save the document to the specified path"""
        os.makedirs(os.path.dirname(filename), exist_ok=True)
        self.document.save(filename)


def generate_report_document(title, sections, output_path, gost_type=None):
    """
    Generate a Word document with multiple sections using AI-generated content
    
    Args:
        title (str): The main title of the document
        sections (list): List of dicts containing section information
        output_path (str): Path where to save the document
        gost_type (str): GOST standard to use ("7.32" or "8.5")
    """
    gost_style = GostStyle(gost_type) if gost_type else None
    generator = WordDocumentGenerator(gost_style)
    generator.add_title(title)

    previous_content = []

    for i, section in enumerate(sections):
        # Create a context-aware prompt by including previous section content
        context_prompt = section["prompt"]
        
        if i > 0 and previous_content:
            # Add context from previous sections
            context = "\n\n".join(previous_content)
            context_prompt = f"""
            Предыдущие разделы содержали следующее содержание:
            ---
            {context}
            ---

            Based on the above content, {section["prompt"]}
            """
        
        # Generate content with context
        content = generator.add_generated_section(
            prompt=context_prompt,
            section_title=section["title"],
            heading_level=section.get("heading_level", 1),
            **(section.get("generation_params", {}))
        )
        
        # Store generated content for context in subsequent sections
        previous_content.append(f"Section {i+1} - {section['title']}: {content}")

    generator.save(output_path)

if __name__ == "__main__":
    sections = [
        {
            "title": "Executive Summary",
            "prompt": "Generate an executive summary for a quarterly sales report",
            "heading_level": 1
        },
        {
            "title": "Market Analysis",
            "prompt": "Generate a market analysis for Q3 2023 focusing on technology sector",
            "heading_level": 1,
            "generation_params": {
                "temperature": 0.7,
                "max_tokens": 1000
            }
        }
    ]
    
    generate_report_document(
        "Quarterly Sales Report - Q3 2023",
        sections,
        "reports/quarterly_report.docx",
        gost_type="7.32"
    )