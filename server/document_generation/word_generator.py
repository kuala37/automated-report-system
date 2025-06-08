from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION, WD_ORIENT
import os
import re
from generation.generate_text_langchain import generate_text, generate_text_with_params

class GostStyle:
    """GOST document styling configuration"""
    def __init__(self, gost_type="8.5"):
        self.gost_type = gost_type
        # ГОСТ 7.32-2017 
        if gost_type == "7.32":
            self.margins = {
                "top": Cm(2), "bottom": Cm(2),
                "left": Cm(2.5), "right": Cm(1.5)
            }
            self.font_name = "Times New Roman"
            self.main_font_size = Pt(14)
            self.heading_font_size = Pt(16)
            self.line_spacing = 1.5
            self.first_line_indent = Cm(1.25)
            self.orientation = "portrait"
            self.page_size = "A4"
        # ГОСТ 8.5 or custom
        else:
            self.margins = {
                "top": Cm(2), "bottom": Cm(2),
                "left": Cm(3), "right": Cm(1.5)
            }
            self.font_name = "Times New Roman"
            self.main_font_size = Pt(12)
            self.heading_font_size = Pt(14)
            self.line_spacing = 1.5
            self.first_line_indent = Cm(1.25)
            self.orientation = "portrait"
            self.page_size = "A4"
            # Дополнительные стили, которые будут заполнены при создании пользовательского стиля
            self.heading_styles = None
            self.list_styles = None

class WordDocumentGenerator:
    def __init__(self, gost_style=None, gost_type=None):
        self.document = Document()
        self.style = gost_style if gost_style else GostStyle()
        self._apply_gost_formatting()
        

    def _apply_gost_formatting(self):
        """Применить форматирование документа на основе выбранного стиля"""
        section = self.document.sections[0]
        
        # Установка полей
        section.top_margin = self.style.margins["top"]
        section.bottom_margin = self.style.margins["bottom"]
        section.left_margin = self.style.margins["left"]
        section.right_margin = self.style.margins["right"]
        
        # Установка ориентации страницы
        if hasattr(self.style, 'orientation') and self.style.orientation == 'landscape':
            section.orientation = WD_ORIENT.LANDSCAPE
        else:
            section.orientation = WD_ORIENT.PORTRAIT

        # Установка стиля абзаца по умолчанию
        style = self.document.styles['Normal']
        font = style.font
        font.name = self.style.font_name
        font.size = self.style.main_font_size
        
        # Установка форматирования абзаца
        paragraph_format = style.paragraph_format
        paragraph_format.line_spacing = self.style.line_spacing
        paragraph_format.first_line_indent = self.style.first_line_indent
        
        self._apply_custom_heading_styles()
        
    def _apply_custom_heading_styles(self):
        """Применить пользовательские стили заголовков, если они определены в объекте стиля"""
        if not hasattr(self.style, 'heading_styles') or not self.style.heading_styles:
            return
            
        heading_styles = self.style.heading_styles
        
        # Словарь для значений выравнивания
        alignment_map = {
            'left': WD_ALIGN_PARAGRAPH.LEFT,
            'center': WD_ALIGN_PARAGRAPH.CENTER,
            'right': WD_ALIGN_PARAGRAPH.RIGHT,
            'justify': WD_ALIGN_PARAGRAPH.JUSTIFY
        }
        
        # Обработка каждого уровня заголовка (h1, h2, h3)
        for level in range(1, 4):
            heading_key = f'h{level}'
            if heading_key in heading_styles:
                h_style = heading_styles[heading_key]
                doc_style = self.document.styles[f'Heading {level}']
                
                # Применить семейство шрифтов
                if 'fontFamily' in h_style:
                    doc_style.font.name = h_style['fontFamily']
                else:
                    doc_style.font.name = self.style.font_name
                
                # Применить размер шрифта
                if 'fontSize' in h_style:
                    doc_style.font.size = Pt(h_style['fontSize'])
                elif level == 1:
                    doc_style.font.size = self.style.heading_font_size
                
                # Применить начертание шрифта
                if 'fontWeight' in h_style:
                    doc_style.font.bold = h_style['fontWeight'] == 'bold'
                    
                # Применить выравнивание текста
                if 'textAlign' in h_style and h_style['textAlign'] in alignment_map:
                    doc_style.paragraph_format.alignment = alignment_map[h_style['textAlign']]
                
                # Применить цвет текста, если указан
                if 'color' in h_style:
                    # Цвет в RGB или именованный цвет, если поддерживается
                    doc_style.font.color.rgb = self._parse_color(h_style['color'])

    def _apply_font_to_heading(self, heading_paragraph, heading_level):
        """Явное применение стиля шрифта к параграфу заголовка"""
        # Проверяем, есть ли кастомные стили для заголовков
        if not hasattr(self.style, 'heading_styles') or not self.style.heading_styles:
            return
        
        heading_key = f'h{heading_level}'
        if heading_key not in self.style.heading_styles:
            return
        
        h_style = self.style.heading_styles[heading_key]
        
        # Получаем все runs в заголовке и применяем стиль к каждому
        for run in heading_paragraph.runs:
            # Шрифт
            if 'fontFamily' in h_style:
                run.font.name = h_style['fontFamily']
            else:
                run.font.name = self.style.font_name
            
            # Размер шрифта
            if 'fontSize' in h_style:
                run.font.size = Pt(float(h_style['fontSize']))
            elif heading_level == 1:
                run.font.size = self.style.heading_font_size
            
            # Стиль шрифта (жирный, курсив и т.д.)
            if 'fontStyle' in h_style:
                font_style = h_style['fontStyle']
                if isinstance(font_style, dict):
                    if 'bold' in font_style:
                        run.font.bold = font_style['bold']
                    if 'italic' in font_style:
                        run.font.italic = font_style['italic']
            elif 'fontWeight' in h_style:
                run.font.bold = h_style['fontWeight'] == 'bold'
            
            # Цвет текста
            if 'color' in h_style:
                run.font.color.rgb = self._parse_color(h_style['color'])
        
        # Для выравнивания (применяется к параграфу целиком)
        if 'textAlign' in h_style:
            alignment_map = {
                'left': WD_ALIGN_PARAGRAPH.LEFT,
                'center': WD_ALIGN_PARAGRAPH.CENTER,
                'right': WD_ALIGN_PARAGRAPH.RIGHT,
                'justify': WD_ALIGN_PARAGRAPH.JUSTIFY
            }
            if h_style['textAlign'] in alignment_map:
                heading_paragraph.alignment = alignment_map[h_style['textAlign']]          

    def _parse_color(self, color_value):
        """Преобразование строки цвета в значение RGB"""
        # Предполагаем, что цвет может быть в формате #RRGGBB
        if color_value.startswith('#') and len(color_value) == 7:
            r = int(color_value[1:3], 16)
            g = int(color_value[3:5], 16)
            b = int(color_value[5:7], 16)
            from docx.shared import RGBColor
            return RGBColor(r, g, b)
        return None

    def add_title(self, title, size=None):
        """Добавляет заголовок документа с указанным форматированием"""
        title_paragraph = self.document.add_paragraph()
        
        # Применяем пользовательский стиль заголовка, если доступен
        if hasattr(self.style, 'heading_styles') and self.style.heading_styles and 'title' in self.style.heading_styles:
            title_style = self.style.heading_styles['title']
            
            # Добавляем текст
            title_run = title_paragraph.add_run(title)
            
            # Применяем шрифт
            if 'fontFamily' in title_style:
                title_run.font.name = title_style['fontFamily']
            else:
                title_run.font.name = self.style.font_name
                
            # Применяем размер шрифта
            if 'fontSize' in title_style:
                title_run.font.size = Pt(title_style['fontSize'])
            else:
                title_run.font.size = self.style.heading_font_size
            
            # Применяем начертание шрифта
            if 'fontWeight' in title_style:
                title_run.font.bold = title_style['fontWeight'] == 'bold'
            
            # Применяем выравнивание
            if 'textAlign' in title_style:
                alignment_map = {
                    'left': WD_ALIGN_PARAGRAPH.LEFT,
                    'center': WD_ALIGN_PARAGRAPH.CENTER,
                    'right': WD_ALIGN_PARAGRAPH.RIGHT,
                    'justify': WD_ALIGN_PARAGRAPH.JUSTIFY
                }
                if title_style['textAlign'] in alignment_map:
                    title_paragraph.alignment = alignment_map[title_style['textAlign']]
        else:
            # Стиль заголовка по умолчанию
            title_run = title_paragraph.add_run(title)
            title_run.font.size = Pt(size) if size else self.style.heading_font_size
            title_run.font.bold = True
            title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
        self.document.add_paragraph()  # Добавляем пустой абзац после заголовка

    def _parse_markdown_content(self, content):
        """Улучшенный парсер markdown контента"""
        lines = content.split('\n')
        parsed_content = []
        current_list = None
        i = 0
        
        while i < len(lines):
            line = lines[i].strip()
            
            # Пропускаем пустые строки
            if not line:
                if current_list:
                    parsed_content.append(current_list)
                    current_list = None
                i += 1
                continue
            
            # Обрабатываем заголовки (должны идти отдельными блоками)
            if line.startswith('####'):
                if current_list:
                    parsed_content.append(current_list)
                    current_list = None
                title = line.replace('###', '').strip()
                parsed_content.append({'type': 'heading', 'level': 4, 'text': title})
                i += 1
                continue

            elif line.startswith('###'):
                if current_list:
                    parsed_content.append(current_list)
                    current_list = None
                title = line.replace('###', '').strip()
                parsed_content.append({'type': 'heading', 'level': 3, 'text': title})
                i += 1
                continue
                
            elif line.startswith('##'):
                if current_list:
                    parsed_content.append(current_list)
                    current_list = None
                title = line.replace('##', '').strip()
                parsed_content.append({'type': 'heading', 'level': 2, 'text': title})
                i += 1
                continue
                
            elif line.startswith('#'):
                if current_list:
                    parsed_content.append(current_list)
                    current_list = None
                title = line.replace('#', '').strip()
                parsed_content.append({'type': 'heading', 'level': 1, 'text': title})
                i += 1
                continue
            
            # Обрабатываем списки
            elif line.startswith(('• ', '- ', '* ')) or re.match(r'^\d+\.', line):
                if not current_list:
                    current_list = {'type': 'list', 'items': [], 'list_type': 'bullet'}
                
                # Определяем тип списка
                if line.startswith(('• ', '- ', '* ')):
                    text = re.sub(r'^[•\-\*]\s*', '', line).strip()
                    current_list['list_type'] = 'bullet'
                else:  # Нумерованный список
                    text = re.sub(r'^\d+\.\s*', '', line).strip()
                    current_list['list_type'] = 'number'
                
                current_list['items'].append(text)
                i += 1
                continue
            
            # Обычный текст - завершаем текущий список и добавляем параграф
            else:
                if current_list:
                    parsed_content.append(current_list)
                    current_list = None
                
                # Собираем многострочный параграф
                paragraph_lines = [line]
                j = i + 1
                
                # Продолжаем читать строки, пока не встретим пустую строку, заголовок или список
                while j < len(lines):
                    next_line = lines[j].strip()
                    
                    # Прерываем на пустой строке, заголовке или списке
                    if (not next_line or 
                        next_line.startswith(('#', '• ', '- ', '* ')) or 
                        re.match(r'^\d+\.', next_line)):
                        break
                    
                    paragraph_lines.append(next_line)
                    j += 1
                
                # Объединяем строки в один параграф
                full_paragraph = ' '.join(paragraph_lines).strip()
                if full_paragraph:
                    parsed_content.append({'type': 'paragraph', 'text': full_paragraph})
                
                i = j
                continue
        
        # Добавляем последний список, если он есть
        if current_list:
            parsed_content.append(current_list)
        
        return parsed_content

    def _parse_inline_formatting(self, text):
        """Улучшенный парсер inline форматирования"""
        if not text:
            return [{'text': '', 'bold': False, 'italic': False}]
        
        parts = []
        current_pos = 0
        
        # Паттерн для поиска форматирования: **жирный** или *курсив*
        pattern = r'(\*\*[^*]+\*\*|\*[^*]+\*)'
        
        for match in re.finditer(pattern, text):
            # Добавляем текст до форматирования
            if match.start() > current_pos:
                plain_text = text[current_pos:match.start()]
                if plain_text:
                    parts.append({'text': plain_text, 'bold': False, 'italic': False})
            
            # Обрабатываем форматированный текст
            matched_text = match.group()
            if matched_text.startswith('**') and matched_text.endswith('**'):
                # Жирный текст
                clean_text = matched_text[2:-2]
                parts.append({'text': clean_text, 'bold': True, 'italic': False})
            elif matched_text.startswith('*') and matched_text.endswith('*'):
                # Курсив (только если это не часть жирного текста)
                clean_text = matched_text[1:-1]
                parts.append({'text': clean_text, 'bold': False, 'italic': True})
            
            current_pos = match.end()
        
        # Добавляем оставшийся текст
        if current_pos < len(text):
            remaining_text = text[current_pos:]
            if remaining_text:
                parts.append({'text': remaining_text, 'bold': False, 'italic': False})
        
        return parts if parts else [{'text': text, 'bold': False, 'italic': False}]

    def add_formatted_content(self, content):
        """Добавляет контент с улучшенной поддержкой markdown"""
        # Предварительная очистка контента
        cleaned_content = self._clean_markdown_content(content)
        parsed_content = self._parse_markdown_content(cleaned_content)
        
        for item in parsed_content:
            if item['type'] == 'heading':
                self._add_formatted_heading(item['text'], item['level'])
            elif item['type'] == 'paragraph':
                if item['text'].strip():  # Пропускаем пустые параграфы
                    self._add_formatted_paragraph(item['text'])
            elif item['type'] == 'list':
                self._add_formatted_list(item['items'], item.get('list_type', 'bullet'))

    def _clean_markdown_content(self, content):
        """Очищает и нормализует markdown контент"""
        # Убираем лишние пробелы и переносы
        lines = content.split('\n')
        cleaned_lines = []
        
        for line in lines:
            # Убираем лишние пробелы в начале и конце
            cleaned_line = line.strip()
            
            # Нормализуем маркеры списков
            if cleaned_line.startswith('•'):
                cleaned_line = '• ' + cleaned_line[1:].strip()
            elif re.match(r'^\d+\.', cleaned_line):
                # Нормализуем нумерацию
                match = re.match(r'^(\d+)\.\s*(.*)', cleaned_line)
                if match:
                    cleaned_line = f"{match.group(1)}. {match.group(2)}"
            
            cleaned_lines.append(cleaned_line)
        
        return '\n'.join(cleaned_lines)

    def _add_formatted_heading(self, text, level):
        """Добавляет заголовок с форматированием"""
        # Ограничиваем уровень заголовков
        level = min(level, 6)
        heading = self.document.add_heading('', level=level)
        
        # Парсим inline форматирование
        formatted_parts = self._parse_inline_formatting(text)
        
        for part in formatted_parts:
            run = heading.add_run(part['text'])
            run.bold = part['bold']
            run.italic = part['italic']
        
        # Применяем кастомное форматирование
        self._apply_font_to_heading(heading, level)

    def _add_formatted_paragraph(self, text):
        """Добавляет параграф с inline форматированием"""
        paragraph = self.document.add_paragraph()
        
        # Парсим inline форматирование
        formatted_parts = self._parse_inline_formatting(text)
        
        for part in formatted_parts:
            run = paragraph.add_run(part['text'])
            run.bold = part['bold']
            run.italic = part['italic']

    def _add_formatted_list(self, items, list_type='bullet'):
        """Добавляет список с форматированием"""
        for item in items:
            if list_type == 'bullet':
                paragraph = self.document.add_paragraph('', style='List Bullet')
            else:
                paragraph = self.document.add_paragraph('', style='List Number')
            
            # Парсим inline форматирование для элемента списка
            formatted_parts = self._parse_inline_formatting(item)
            
            for part in formatted_parts:
                run = paragraph.add_run(part['text'])
                run.bold = part['bold']
                run.italic = part['italic']
    
    def add_section(self, title, content, heading_level=1):
        """Добавить раздел с заголовком и содержимым"""
        # Добавляем заголовок с соответствующим уровнем
        heading = self.document.add_heading(title, level=heading_level)
        self._apply_font_to_heading(heading, heading_level)

        # Добавляем содержимое с поддержкой markdown
        self.add_formatted_content(content)

    def add_generated_section(self, prompt, section_title, heading_level=1, **generation_params):
        """Добавить раздел с автоматически сгенерированным содержимым"""
        # Генерация контента с использованием ИИ сервиса
        try:
            content = generate_text_with_params(prompt, **generation_params)
        except Exception as e:
            print(f"Ошибка генерации контента: {str(e)}")
            content = f"[Ошибка генерации контента: {str(e)}]"
        
        # Добавляем заголовок раздела
        heading = self.document.add_heading(section_title, level=heading_level)
        self._apply_font_to_heading(heading, heading_level)
        
        # Очищаем и нормализуем контент перед добавлением
        print(f"🔍 Исходный контент от ИИ ({len(content)} символов):")
        print(f"'{content[:200]}...'")
        
        # Добавляем сгенерированный контент с улучшенным форматированием
        self.add_formatted_content(content)
        
        print(f"✅ Контент добавлен в документ")
        return content
    
    def add_paragraph_text(self, text):
        """Добавить простой абзац текста со стилем"""
        paragraph = self.document.add_paragraph()
        
        # Добавить текст в абзац
        paragraph.add_run(text)
        
        # Применить стиль абзаца, если он определен в пользовательском стиле
        if hasattr(self.style, 'paragraphs') and self.style.paragraphs:
            p_style = self.style.paragraphs
            
            # Применить выравнивание текста, если указано
            if 'textAlign' in p_style:
                alignment_map = {
                    'left': WD_ALIGN_PARAGRAPH.LEFT,
                    'center': WD_ALIGN_PARAGRAPH.CENTER,
                    'right': WD_ALIGN_PARAGRAPH.RIGHT,
                    'justify': WD_ALIGN_PARAGRAPH.JUSTIFY
                }
                if p_style['textAlign'] in alignment_map:
                    paragraph.alignment = alignment_map[p_style['textAlign']]
               
    def save(self, filename):
        """Сохранить документ по указанному пути"""
        self.document.save(filename)

    # Добавление методов для анализа и выполнения команд редактирования
    def update_paragraph_text(self, paragraph_index, new_text):
        """Обновляет текст конкретного параграфа"""
        if 0 <= paragraph_index < len(self.document.paragraphs):
            paragraph = self.document.paragraphs[paragraph_index]
            # Очищаем существующие runs
            for i in range(len(paragraph.runs)):
                paragraph.runs[0]._element.getparent().remove(paragraph.runs[0]._element)
            
            # Добавляем новый run с обновленным текстом
            run = paragraph.add_run(new_text)
            return True
        return False

    def replace_text(self, old_text, new_text):
        """Заменяет все вхождения текста в документе"""
        found = False
        for paragraph in self.document.paragraphs:
            if old_text in paragraph.text:
                for run in paragraph.runs:
                    run.text = run.text.replace(old_text, new_text)
                found = True
        return found

    def format_text_in_paragraph(self, paragraph_index, start, end, formatting):
        """Применяет форматирование к части текста в параграфе"""
        if 0 <= paragraph_index < len(self.document.paragraphs):
            paragraph = self.document.paragraphs[paragraph_index]
            text = paragraph.text
            
            if start < 0 or end > len(text) or start >= end:
                return False
                
            # Сложная логика для форматирования части текста
            # ...
            
            return True
        return False