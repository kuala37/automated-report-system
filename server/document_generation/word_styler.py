from docx import Document
from docx.shared import Pt, RGBColor
import re

def hex_to_rgb(hex_color):
    """Преобразует цвет из HEX формата в RGB."""
    hex_color = hex_color.lstrip('#')
    return tuple(int(hex_color[i:i+2], 16) for i in (0, 2, 4))

def apply_styles_to_word(doc_path, content, styles):
    """
    Применяет выбранные стили к Word документу.
    
    Args:
        doc_path: Путь для сохранения документа
        content: Текстовое содержимое документа
        styles: Словарь со стилями
    
    Returns:
        Путь к сохраненному документу
    """
    # Создаем новый документ
    doc = Document()
    
    # Настройка стилей
    style_normal = doc.styles['Normal']
    style_normal.font.name = styles.get('font_main', 'Calibri')
    style_normal.font.size = Pt(styles.get('font_size_main', 11))
    
    # Настройка стилей заголовков
    for i, heading_style in enumerate(['Heading 1', 'Heading 2', 'Heading 3'], 1):
        if heading_style in doc.styles:
            style = doc.styles[heading_style]
            style.font.name = styles.get('font_headings', 'Arial')
            if i == 1:
                style.font.size = Pt(styles.get('font_size_h1', 16))
            elif i == 2:
                style.font.size = Pt(styles.get('font_size_h2', 14))
            else:
                style.font.size = Pt(styles.get('font_size_h2', 14) - 2)
                
            # Преобразуем HEX в RGB для заголовка
            rgb = hex_to_rgb(styles.get('color_headings', '#1F497D'))
            style.font.color.rgb = RGBColor(rgb[0], rgb[1], rgb[2])
    
    # Парсинг контента и добавление в документ
    lines = content.split('\n')
    current_heading_level = 0
    
    for line in lines:
        # Проверка на заголовок по символам '#' (как в markdown)
        if line.strip().startswith('#'):
            heading_count = len(re.match(r'^#+', line.strip()).group())
            if heading_count <= 3:  # Поддерживаем до 3 уровней заголовков
                heading_text = line.strip('#').strip()
                if heading_count == 1:
                    doc.add_heading(heading_text, level=1)
                else:
                    doc.add_heading(heading_text, level=heading_count)
                current_heading_level = heading_count
            else:
                doc.add_paragraph(line)
        # Проверка на заголовок по другим признакам
        elif line.strip() and not line.strip()[0].islower() and len(line.strip()) < 100:
            # Эвристика: короткие строки, начинающиеся с заглавной буквы, 
            # могут быть заголовками
            if current_heading_level == 0:
                doc.add_heading(line.strip(), level=1)
                current_heading_level = 1
            else:
                doc.add_heading(line.strip(), level=current_heading_level + 1)
        else:
            if line.strip():
                doc.add_paragraph(line)
    
    # Сохранение документа
    doc.save(doc_path)
    return doc_path