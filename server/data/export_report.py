from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from docx import Document

def export_to_pdf(text, filename):
    """
    Экспортирует текст в PDF-файл.

    Аргументы:
        text (str): Текст для экспорта.
        filename (str): Имя файла (без расширения).
    """
    pdf_path = f"{filename}.pdf"
    c = canvas.Canvas(pdf_path, pagesize=letter)
    c.drawString(72, 750, "Сгенерированный отчет")
    c.drawString(72, 730, "-" * 50)
    text_lines = text.split("\n")
    y = 700
    for line in text_lines:
        c.drawString(72, y, line)
        y -= 12
    c.save()
    return pdf_path

def export_to_docx(text, filename):
    """
    Экспортирует текст в DOCX-файл.

    Аргументы:
        text (str): Текст для экспорта.
        filename (str): Имя файла (без расширения).
    """
    docx_path = f"{filename}.docx"
    doc = Document()
    doc.add_heading("Сгенерированный отчет", 0)
    doc.add_paragraph(text)
    doc.save(docx_path)
    return docx_path

def export_to_html(text, filename):
    """
    Экспортирует текст в HTML-файл.

    Аргументы:
        text (str): Текст для экспорта.
        filename (str): Имя файла (без расширения).
    """
    html_path = f"{filename}.html"
    with open(html_path, "w", encoding="utf-8") as f:
        f.write(f"<h1>Сгенерированный отчет</h1>\n<p>{text}</p>")
    return html_path

# Пример использования
if __name__ == "__main__":
    try:
        # Пример текста отчета
        report_text = (
            "Отчет о продажах за последний квартал.\n"
            "Общий объем продаж: 1 200 000 руб.\n"
            "Основные тренды: рост продаж в категории электроники.\n"
            "Рекомендации: продолжить использовать email-рассылки."
        )

        pdf_file = export_to_pdf(report_text, "report")
        print(f"PDF-файл создан: {pdf_file}")

        docx_file = export_to_docx(report_text, "report")
        print(f"DOCX-файл создан: {docx_file}")

        html_file = export_to_html(report_text, "report")
        print(f"HTML-файл создан: {html_file}")

    except Exception as e:
        print(e)