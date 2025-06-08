import json
from typing import Dict, Any, Optional
from sqlalchemy.ext.asyncio import AsyncSession
from models.models import Report
from generation.generate_text_langchain import generate_text_with_params
from services.document_editor_service import DocumentEditorService
import re


class SmartDocumentAgent:
    """Умный агент для редактирования документов через прямые LLM вызовы"""
    
    def __init__(self):
        self.editor_service = DocumentEditorService()
    
    async def process_command(
        self, 
        db: AsyncSession, 
        report_id: int, 
        user_id: int, 
        command_text: str
    ) -> Dict[str, Any]:
        """Обрабатывает команду пользователя напрямую через LLM"""
        
        print(f"🤖 SmartAgent получил команду: '{command_text}'")
        
        try:
            # Получаем отчет и контекст
            report = await db.get(Report, report_id)
            if not report or report.user_id != user_id:
                return {"success": False, "message": "Отчет не найден или нет доступа"}
            
            # Получаем превью документа
            document_text = await self._get_document_text(report)
            print(f"📄 Получен текст документа: {len(document_text)} символов")
            
            # Этап 1: Умный анализ команды через LLM
            analysis_result = await self._analyze_command_with_llm(command_text, document_text)
            print(f"🧠 Анализ LLM: {analysis_result}")
            
            if not analysis_result["success"]:
                return analysis_result
            
            # Этап 2: Выполнение действия на основе анализа
            action_result = await self._execute_action(
                db, report_id, analysis_result["action"], user_id
            )
            print(f"⚡ Результат выполнения: {action_result}")
            
            return action_result
            
        except Exception as e:
            print(f"❌ Ошибка в SmartAgent: {str(e)}")
            return {"success": False, "message": f"Ошибка при обработке команды: {str(e)}"}
    
    async def _analyze_command_with_llm(self, command_text: str, document_text: str) -> Dict[str, Any]:
        """Анализирует команду через LLM и определяет действие"""
        selected_text = None
        selected_paragraph_id = None
        clean_command = command_text

        match = re.search(r'\[ВЫДЕЛЕННЫЙ ТЕКСТ: "([^"]*(?:"[^"]*"[^"]*)*)"(?:\s+в параграфе (\d+))?\]\s*(.*)', command_text, re.DOTALL)
        # match = re.search(r'\[ВЫДЕЛЕННЫЙ ТЕКСТ: "([^"]+)"(?:\s+в параграфе (\d+))?\]\s*(.*)', command_text)
        
        if match:
            selected_text = match.group(1)
            selected_paragraph_id = match.group(2)
            clean_command = match.group(3)
            print(f"🎯 Найден выделенный текст: '{selected_text}', команда: '{clean_command}'")
        
        # Ограничиваем текст документа для экономии токенов
        doc_preview = document_text[:1500] if len(document_text) > 1500 else document_text
        
        analysis_prompt = f"""
Ты - эксперт по редактированию документов. Пользователь дал команду для редактирования документа.

КОМАНДА ПОЛЬЗОВАТЕЛЯ: "{clean_command}"

{f'ВЫДЕЛЕННЫЙ ПОЛЬЗОВАТЕЛЕМ ТЕКСТ: "{selected_text}"' if selected_text else 'ВЫДЕЛЕННЫЙ ТЕКСТ: отсутствует'}
{f'ПАРАГРАФ: {selected_paragraph_id}' if selected_paragraph_id else 'ПАРАГРАФ: не указан'}

ТЕКСТ ДОКУМЕНТА:
{doc_preview}

Твоя задача - понять, что именно хочет сделать пользователь, и дать точные инструкции.

КРИТИЧЕСКИ ВАЖНЫЕ ПРАВИЛА:
-ЕСЛИ ЕСТЬ ВЫДЕЛЕННЫЙ ТЕКСТ - используй его ТОЧНО КАК ЕСТЬ в поле "target"! 
-НЕ ИЗМЕНЯЙ, НЕ ДОПОЛНЯЙ и НЕ СОКРАЩАЙ выделенный текст!
-target должен ПОЛНОСТЬЮ соответствовать выделенному тексту!

ВАЖНЫЕ ПРАВИЛА:

1. Если пользователь говорит "перефразируй текст" БЕЗ указания конкретного текста - он имеет в виду весь документ или первый абзац
2. Если говорит "перефразируй первый абзац" - найди первый абзац в документе
3. Если просит сделать что-то жирным/курсивом - это форматирование
4. Если просит заменить конкретные слова - это замена
5. Если команда "удали параграф" и есть выделенный текст с paragraph_id, используй этот ID
6. Если команда содержит "первый", "второй" и т.д., извлеки номер параграфа ввиде обычной цифры (1,2,3 и т.д.) и используй его как paragraph_id
7. Если команда "удали этот параграф" без указания номера, но есть выделение - используй paragraph_id из выделения
8. Если пользователь не просит КОНКРЕТНО удалить параграф,то используй используй "удали текст" с выделенным текстом
9. Если команда неясна, верни "clarify" с объяснением, что нужно уточнить

ВОЗМОЖНЫЕ ДЕЙСТВИЯ:
- rewrite_all: перефразировать весь документ
- rewrite_paragraph: перефразировать конкретный абзац (нужно указать какой)
- replace_text: заменить один текст на другой
- remove_formatting: снятие конкретного форматирования
- remove_all_formatting: снятие всего форматирования
- format_text: сделать текст жирным/курсивом/подчеркнутым
- add_text: добавить новый текст
- add_paragraph: добавить новый параграф
- delete_text: удалить текст
- delete_paragraph: удалить параграф
- format_all_text: форматирование ВСЕГО документа
- format_all_headings: форматирование всех заголовков


Проанализируй команду и верни JSON в формате:
{{
    "action": "название_действия",
    "target": "конкретный текст для обработки",
    "replacement": "новый текст (если нужна замена)",
    "style": "bold/italic/underline (если форматирование)",
    "explanation": "объяснение что будем делать"
}}

Если команда неясна, верни:
{{
    "action": "clarify",
    "explanation": "Объяснение что именно нужно уточнить"
}}

Отвечай ТОЛЬКО JSON без дополнительного текста.
"""
        
        try:
            response = generate_text_with_params(
                analysis_prompt, 
                temperature=0.1, 
                max_tokens=400
            )
            
            print(f"🔍 Ответ LLM на анализ: {response}")
            
            # Извлекаем JSON из ответа
            response_clean = response.strip()
            if response_clean.startswith('```json'):
                response_clean = response_clean[7:-3]
            elif response_clean.startswith('```'):
                response_clean = response_clean[3:-3]
            
            try:
                analysis = json.loads(response_clean)
                
                # ВАЖНО: Если target не указан, но есть выделенный текст - используем его
                # if not analysis.get("target") and selected_text:
                #     analysis["target"] = 
                print("Выделенный текст:", selected_text)
                if selected_text:
                    print(f"🔒 ПРИНУДИТЕЛЬНО заменяем target '{analysis.get('target', 'НЕТ')}' на выделенный текст: '{selected_text}'")
                    analysis["target"] = selected_text
                    
                # Добавляем paragraph_id если есть
                if selected_paragraph_id and not analysis.get("paragraph_id"):
                    analysis["paragraph_id"] = int(selected_paragraph_id)
                else:
                    print(f"⚠️ Выделенный текст не найден, используем target от LLM: '{analysis.get('target', 'НЕТ')}'")
            
                return {"success": True, "action": analysis}
                
            except json.JSONDecodeError:
                # Если JSON невалидный, пытаемся найти его в тексте
                json_match = re.search(r'\{[^{}]*\}', response, re.DOTALL)
                if json_match:
                    analysis = json.loads(json_match.group(0))
                    if not analysis.get("target") and selected_text:
                        analysis["target"] = selected_text
                    if selected_paragraph_id and not analysis.get("paragraph_id"):
                        analysis["paragraph_id"] = int(selected_paragraph_id)
                    return {"success": True, "action": analysis}
                else:
                    return {"success": False, "message": "LLM не смог распознать команду"}
                    
        except Exception as e:
            print(f"❌ Ошибка при анализе через LLM: {str(e)}")
            return {"success": False, "message": f"Ошибка анализа команды: {str(e)}"}


    async def _execute_action(
        self, 
        db: AsyncSession, 
        report_id: int, 
        action_data: Dict[str, Any], 
        user_id: int
    ) -> Dict[str, Any]:
        """Выполняет действие на основе результатов анализа LLM"""
        
        action_type = action_data.get("action")
        
        if action_type == "clarify":
            return {
                "success": False, 
                "message": f"Нужно уточнение: {action_data.get('explanation', 'Команда неясна')}"
            }
        
        elif action_type == "rewrite_all":
            return await self._rewrite_entire_document(db, report_id, user_id)
        
        elif action_type == "rewrite_paragraph":
            target_text = action_data.get("target", "")
            return await self._rewrite_specific_text(db, report_id, target_text, user_id)
        
        elif action_type == "replace_text":
            old_text = action_data.get("target", "")
            new_text = action_data.get("replacement", "")
            return await self._replace_text_in_document(db, report_id, old_text, new_text, user_id)
        
        elif action_type == "format_text":
            target_text = action_data.get("target", "")
            style = action_data.get("style", "bold")
            return await self._format_text_in_document(db, report_id, target_text, style, user_id)

        elif action_type == "format_all_text":
            style = action_data.get("style", "bold")
            return await self._format_all_text_in_document(db, report_id, style, user_id)
        
        elif action_type == "format_all_headings":
            style = action_data.get("style", "bold")
            return await self._format_all_headings_in_document(db, report_id, style, user_id)

        elif action_type == "add_text":
            new_text = action_data.get("replacement", "")
            return await self._add_text_to_document(db, report_id, new_text, user_id)
        
        elif action_type == "delete_text":
            target_text = action_data.get("target", "")
            return await self._delete_text_from_document(db, report_id, target_text, user_id)
        
        elif action_type == "delete_paragraph":
            paragraph_id = action_data.get("paragraph_id")
            if not paragraph_id:
                # Если не указан ID параграфа, но есть выделенный текст, пытаемся его найти
                paragraph_id = action_data.get("target", "")

            if paragraph_id is None:
                return {"success": False, "message": "Не удалось определить какой параграф удалить. Выделите текст в нужном параграфе."}
            return await self._delete_paragraph_from_document(db, report_id, paragraph_id, user_id)
        

        elif action_type == "remove_formatting":
            paragraph_id = action_data.get("paragraph_id")
            target_text = action_data.get("target", "")
            style = action_data.get("style", "bold")
            return await self._remove_formatting_from_text(db, report_id, target_text, style,paragraph_id, user_id)
        
        elif action_type == "remove_all_formatting":
            target_text = action_data.get("target", "")
            return await self._remove_all_formatting_from_text(db, report_id, target_text, user_id)
        
    async def _rewrite_entire_document(self, db: AsyncSession, report_id: int, user_id: int) -> Dict[str, Any]:
        """Перефразирует весь документ"""
        
        # Получаем текст документа
        report = await db.get(Report, report_id)
        document_text = await self._get_document_text(report)
        
        if len(document_text) < 50:
            return {"success": False, "message": "Документ слишком короткий для перефразирования"}
        
        # Генерируем улучшенную версию через LLM
        rewrite_prompt = f"""
Перефразируй следующий документ, сохраняя его смысл, но улучшив стиль, ясность и читаемость:

{document_text}

Требования:
- Сохрани структуру документа
- Улучши стиль изложения
- Сделай текст более ясным и профессиональным
- Исправь грамматические ошибки
- Сохрани все ключевые идеи и факты

Верни только переписанный текст без дополнительных комментариев.
"""
        
        try:
            new_text = generate_text_with_params(
                rewrite_prompt,
                temperature=0.7,
                max_tokens=2000
            )
            
            # Заменяем весь документ
            edit_command = {
                "command": "replace_all_content",
                "newText": new_text,
                "user_id": user_id
            }
            
            result = await self.editor_service.update_document_with_edit(db, report_id, edit_command)
            return result
            
        except Exception as e:
            return {"success": False, "message": f"Ошибка при перефразировании: {str(e)}"}
    
    async def _rewrite_specific_text(self, db: AsyncSession, report_id: int, target_text: str, user_id: int) -> Dict[str, Any]:
        """Перефразирует конкретный текст"""
        
        if not target_text or len(target_text) < 10:
            # Если целевой текст не указан, берем первый абзац
            report = await db.get(Report, report_id)
            document_text = await self._get_document_text(report)
            paragraphs = document_text.split('\n\n')
            target_text = paragraphs[0] if paragraphs else document_text[:200]
        
        rewrite_prompt = f"""
Перефразируй следующий текст, сохранив его смысл, но улучшив стиль и ясность:

"{target_text}"

Верни только переписанный текст без кавычек и дополнительных комментариев.
"""
        
        try:
            new_text = generate_text_with_params(
                rewrite_prompt,
                temperature=0.7,
                max_tokens=500
            )
            
            # Заменяем текст в документе
            edit_command = {
                "command": "replace_text",
                "oldText": target_text,
                "newText": new_text.strip(),
                "user_id": user_id
            }
            
            result = await self.editor_service.update_document_with_edit(db, report_id, edit_command)
            return result
            
        except Exception as e:
            return {"success": False, "message": f"Ошибка при перефразировании текста: {str(e)}"}
    
    async def _replace_text_in_document(self, db: AsyncSession, report_id: int, old_text: str, new_text: str, user_id: int) -> Dict[str, Any]:
        """Заменяет текст в документе"""
        
        edit_command = {
            "command": "replace_text",
            "oldText": old_text,
            "newText": new_text,
            "user_id": user_id
        }
        
        return await self.editor_service.update_document_with_edit(db, report_id, edit_command)
    
    async def _format_text_in_document(self, db: AsyncSession, report_id: int, target_text: str, style: str, user_id: int) -> Dict[str, Any]:
        """Форматирует текст в документе"""
        
        edit_command = {
            "command": "format_text",
            "text": target_text,
            "style": style,
            "user_id": user_id
        }
        
        return await self.editor_service.update_document_with_edit(db, report_id, edit_command)

    async def _format_all_text_in_document(self, db: AsyncSession, report_id: int, style: str, user_id: int) -> Dict[str, Any]:
        """Форматирует весь текст в документе"""
        
        edit_command = {
            "command": "format_all_text",
            "style": style,
            "user_id": user_id
        }
        
        return await self.editor_service.update_document_with_edit(db, report_id, edit_command)
    
    async def _format_all_headings_in_document(self, db: AsyncSession, report_id: int, style: str, user_id: int) -> Dict[str, Any]:
        """Форматирует все заголовки в документе"""
        
        edit_command = {
            "command": "format_all_headings",
            "style": style,
            "user_id": user_id
        }
        
        return await self.editor_service.update_document_with_edit(db, report_id, edit_command)

    async def _add_text_to_document(self, db: AsyncSession, report_id: int, new_text: str, user_id: int) -> Dict[str, Any]:
        """Добавляет текст в документ"""
        
        edit_command = {
            "command": "add_paragraph",
            "text": new_text,
            "user_id": user_id
        }
        
        return await self.editor_service.update_document_with_edit(db, report_id, edit_command)
    
    async def _delete_text_from_document(self, db: AsyncSession, report_id: int, target_text: str, user_id: int) -> Dict[str, Any]:
        """Удаляет текст из документа"""
        
        edit_command = {
            "command": "replace_text",
            "oldText": target_text,
            "newText": "",
            "user_id": user_id
        }
        
        return await self.editor_service.update_document_with_edit(db, report_id, edit_command)
    
    async def _delete_paragraph_from_document(self, db: AsyncSession, report_id: int, paragraph_id: int, user_id: int) -> Dict[str, Any]:
        """Удаляет параграф из документа"""
        
        edit_command = {
            "command": "delete_paragraph",
            "paragraphId": paragraph_id,
            "user_id": user_id
        }
        
        return await self.editor_service.update_document_with_edit(db, report_id, edit_command)
        
    async def _remove_formatting_from_text(self, db: AsyncSession, report_id: int, target_text: str, style: str, paragraph_id: int,user_id: int) -> Dict[str, Any]:
        """Снимает конкретное форматирование с текста"""
        
        edit_command = {
            "command": "remove_formatting",
            "text": target_text,
            "style": style,
            "paragraphId": paragraph_id,
            "user_id": user_id
        }
        
        return await self.editor_service.update_document_with_edit(db, report_id, edit_command)

    async def _remove_all_formatting_from_text(self, db: AsyncSession, report_id: int, target_text: str, user_id: int) -> Dict[str, Any]:
        """Снимает все форматирование с текста"""
        
        edit_command = {
            "command": "remove_all_formatting",
            "text": target_text,
            "user_id": user_id
        }
        
        return await self.editor_service.update_document_with_edit(db, report_id, edit_command)



    async def _get_document_text(self, report: Report) -> str:
        """Получает текст документа"""
        
        if hasattr(report, 'html_content') and report.html_content:
            try:
                from bs4 import BeautifulSoup
                soup = BeautifulSoup(report.html_content, 'html.parser')
                text = soup.get_text()
                # Очищаем и форматируем текст
                lines = [line.strip() for line in text.split('\n') if line.strip()]
                return '\n\n'.join(lines)
            except Exception as e:
                print(f"Ошибка извлечения текста из HTML: {str(e)}")
        
        # Fallback: читаем из docx файла
        try:
            if hasattr(report, 'file_path') and report.file_path:
                from docx import Document
                doc = Document(report.file_path)
                paragraphs = []
                for para in doc.paragraphs:
                    if para.text.strip():
                        paragraphs.append(para.text.strip())
                return '\n\n'.join(paragraphs)
        except Exception as e:
            print(f"Ошибка чтения docx файла: {str(e)}")
        
        return "Не удалось получить текст документа"